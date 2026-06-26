"""
Exam Material PDF Generator
============================
Extracts content from source PDFs/DOCX, processes each chapter
through GPT-4o-mini to generate structured exam notes, then
compiles into professional PDFs per class+medium.

Output: 6 PDFs (3 classes x 2 mediums)
"""

import os
import sys
import io
import re
import time
import json
import traceback
from pathlib import Path
from collections import OrderedDict

# PDF/DOCX reading
from PyPDF2 import PdfReader
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

# PDF writing
from fpdf import FPDF

# OpenAI
import openai

# Groq from llm_client
import sys
sys.path.append(str(Path(__file__).resolve().parent))
from exam_notes_generator.core.llm_client import generate_with_llm

# ─── Configuration ─────────────────────────────────────────────────────────────
BASE_DIR = Path(r"C:\Users\bhask\Desktop\New folder")
OUTPUT_DIR = BASE_DIR / "Generated_Exam_PDFs"
OUTPUT_DIR.mkdir(exist_ok=True)

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")

# Force UTF-8 output on Windows console
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
client = None  # initialized in main()

ENGLISH_MEDIUM = {
    "6th": BASE_DIR / "English" / "6th class",
    "7th": BASE_DIR / "English" / "7th class",
    "8th": BASE_DIR / "English" / "8th class",
}
HINDI_MEDIUM = {
    "6th": BASE_DIR / "Hindi" / "Class 6th",
    "7th": BASE_DIR / "Hindi" / "Class 7th",
    "8th": BASE_DIR / "Hindi" / "Class 8th",
    "11th_Commerce": BASE_DIR / "Hindi" / "11th &12th" / "Class 11th Commerce",
    "11th_Humanities": BASE_DIR / "Hindi" / "11th &12th" / "Class 11th Humanities",
    "11th_Science": BASE_DIR / "Hindi" / "11th &12th" / "Class 11th Science",
    "12th_Commerce": BASE_DIR / "Hindi" / "11th &12th" / "Class 12th Commerce",
    "12th_Humanities": BASE_DIR / "Hindi" / "11th &12th" / "Class 12th Humanities",
    "12th_Science": BASE_DIR / "Hindi" / "11th &12th" / "Class 12th Science",
}

# ─── GPT Prompt (matches user's exact format) ─────────────────────────────────

SYSTEM_PROMPT_BASE = """You are an expert CBSE teacher creating concise, exam-oriented revision notes.

Generate structured notes in EXACTLY this format. Do NOT add extra sections.
Keep the section headings CLEAN — do NOT include any instructions or descriptions in the headings.

## 1. Introduction
Write 2-3 sentences about what this chapter is about and why it matters.

## 2. Key Concepts
List 5-8 most important concepts as bullet points.

## 3. Important Formulas & Key Terms
- For Science/Maths: List all important formulas with symbols
- For English/Hindi: List 3-4 literary devices with examples from this chapter
- For Social Science: List key definitions and important dates/facts

## 4. Important Exam Points
List 8-10 high-priority bullet points likely to be asked in exams.

## 5. Important Questions
List 5 important questions.

## 6. Quick Summary
Write a 3 sentence quick revision summary.

CRITICAL — NO LATEX ALLOWED:
* NEVER use dollar signs ($) around any formula or value.
* NEVER use LaTeX syntax like \\frac{}{}, \\sqrt{}, \\times, \\cdot, etc.
* Use plain text for formulas: a^2 + b^2 = c^2, Speed = Distance / Time
* Plain readable text ONLY."""


def get_system_prompt_for_medium(medium):
    """Return system prompt with language instruction based on medium."""
    if medium == "Hindi":
        return SYSTEM_PROMPT_BASE + "\n\nLANGUAGE RULE: You MUST write the ENTIRE response in Hindi (Devanagari script). All headings, explanations, bullet points, questions, and summaries must be in Hindi. Do NOT write in English."
    else:
        return SYSTEM_PROMPT_BASE + "\n\nLANGUAGE RULE: Write the entire response in English."


# ─── Text Extraction ──────────────────────────────────────────────────────────

def extract_text_from_pdf(pdf_path):
    try:
        reader = PdfReader(str(pdf_path))
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text.strip()
    except Exception as e:
        print(f"    [WARN] PDF: {pdf_path.name} - {e}")
        return ""

def extract_text_from_docx(docx_path):
    try:
        doc = Document(str(docx_path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        print(f"    [WARN] DOCX: {docx_path.name} - {e}")
        return ""

def extract_text(file_path):
    s = file_path.suffix.lower()
    if s == ".pdf":
        return extract_text_from_pdf(file_path)
    elif s == ".docx":
        return extract_text_from_docx(file_path)
    return ""

def clean_chapter_name(filename):
    name = filename.stem
    name = re.sub(r'_create$', '', name, flags=re.IGNORECASE)
    name = name.replace('_', ' ')
    name = re.sub(r'\s+', ' ', name).strip()
    name = re.sub(r'^(Chapter|Ch|CH)\s*[-_.]\s*', 'Chapter ', name, flags=re.IGNORECASE)
    name = re.sub(r'^(Chapter|Ch|CH)\s+(\d+)\s+', r'Chapter \2: ', name, flags=re.IGNORECASE)
    return name

def sort_key(name):
    m = re.search(r'(\d+)', name)
    return int(m.group(1)) if m else 999


# ─── Collect Files ─────────────────────────────────────────────────────────────

def collect_subject_data(class_dir):
    subjects = OrderedDict()
    if not class_dir.exists():
        return subjects

    for subject_dir in sorted(class_dir.iterdir()):
        if not subject_dir.is_dir():
            continue
        subject_name = subject_dir.name
        sub_categories = OrderedDict()

        has_subcats = any(
            d.is_dir() and (list(d.rglob("*.pdf")) or list(d.rglob("*.docx")))
            for d in subject_dir.iterdir() if d.is_dir()
        )

        def _gather(folder):
            files = sorted(list(folder.rglob("*.pdf")) + list(folder.rglob("*.docx")))
            seen = {}
            for f in files:
                stem = re.sub(r'_create$', '', f.stem, flags=re.IGNORECASE)
                if stem not in seen or f.suffix.lower() == '.pdf':
                    seen[stem] = f
            chapters = []
            for f in sorted(seen.values(), key=lambda x: sort_key(x.name)):
                text = extract_text(f)
                if text and len(text) > 30:
                    chapters.append({"name": clean_chapter_name(f), "text": text})
            return chapters

        if has_subcats:
            for sub_dir in sorted(subject_dir.iterdir()):
                if sub_dir.is_dir():
                    chs = _gather(sub_dir)
                    if chs:
                        sub_categories[sub_dir.name] = chs
        else:
            chs = _gather(subject_dir)
            if chs:
                sub_categories["_main"] = chs

        if sub_categories:
            subjects[subject_name] = sub_categories
            total = sum(len(c) for c in sub_categories.values())
            subs = ", ".join(f"{k}({len(v)})" for k, v in sub_categories.items() if k != "_main")
            print(f"    {subject_name}: {total} ch" + (f" [{subs}]" if subs else ""))

    return subjects


# ─── GPT Processing ───────────────────────────────────────────────────────────

def process_chapter_gpt(chapter_name, chapter_text, subject, class_name, medium="English"):
    """Send chapter text to GPT and get structured exam notes."""
    global client
    # Truncate very long chapters to stay within token limits
    max_chars = 10000
    if len(chapter_text) > max_chars:
        chapter_text = chapter_text[:max_chars] + "\n\n[Content truncated for processing...]"
    
    system_prompt = get_system_prompt_for_medium(medium)
    
    user_msg = f"""Subject: {subject}
Class: {class_name}
Chapter: {chapter_name}
Language: {medium}

---
SOURCE CONTENT:
{chapter_text}
---

Generate structured exam short notes for this chapter in the specified format."""

    try:
        # Select model based on API provider
        if OPENAI_API_KEY.lower() == "groq":
            return generate_with_llm(
                system_prompt=system_prompt,
                user_prompt=user_msg,
                temperature=0.25,
                max_tokens=2500,
                context_length=8192
            )
        
        if OPENAI_API_KEY.startswith("sk-or-"):
            model = "openai/gpt-4o-mini"
        elif OPENAI_API_KEY.startswith("AIza"):
            model = "gemini-2.0-flash"
        elif OPENAI_API_KEY.lower() == "ollama":
            model = "qwen2.5:7b" # Using the model you are pulling
        else:
            model = "gpt-4o-mini"
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_msg}
            ],
            temperature=0.25,
            max_tokens=2500,
        )
        return resp.choices[0].message.content
    except openai.RateLimitError:
        print(f"      [RATE LIMIT] Waiting 30s...")
        time.sleep(30)
        try:
            if OPENAI_API_KEY.startswith("sk-or-"):
                model = "openai/gpt-4o-mini"
            elif OPENAI_API_KEY.startswith("AIza"):
                model = "gemini-2.0-flash"
            elif OPENAI_API_KEY.lower() == "ollama":
                model = "qwen2.5:7b"
            else:
                model = "gpt-4o-mini"
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_msg}
                ],
                temperature=0.25,
                max_tokens=2500,
            )
            return resp.choices[0].message.content
        except Exception as e2:
            print(f"      [ERROR] Retry failed: {e2}")
            return None
    except Exception as e:
        print(f"      [ERROR] GPT failed: {e}")
        return None


# ─── PDF Generation ───────────────────────────────────────────────────────────

class ExamPDF(FPDF):
    def __init__(self, class_name, medium):
        super().__init__()
        self.class_name = class_name
        self.medium = medium
        self._subj = ""
        self.set_auto_page_break(auto=True, margin=20)
        # Register Unicode fonts supporting both Latin and Devanagari (Hindi)
        _core_dir = Path(r"C:\Users\bhask\Desktop\New folder\exam_notes_generator\core")
        _deva_regular = _core_dir / "NotoSansDevanagari-Regular.ttf"
        _deva_bold = _core_dir / "NotoSansDevanagari-Bold.ttf"
        self.add_font("Body", "", "C:/Windows/Fonts/arial.ttf")
        self.add_font("Body", "B", "C:/Windows/Fonts/arialbd.ttf")
        self.add_font("Body", "I", "C:/Windows/Fonts/ariali.ttf")
        self.add_font("Body", "BI", "C:/Windows/Fonts/arialbi.ttf")
        if _deva_regular.exists() and _deva_bold.exists():
            self.add_font("Devanagari", "", str(_deva_regular))
            self.add_font("Devanagari", "B", str(_deva_bold))
            self.set_fallback_fonts(["Devanagari"])
        self.set_text_shaping(True)

    def header(self):
        if self.page_no() <= 1:
            return
        self.set_font('Body', 'I', 8)
        self.set_text_color(130, 130, 130)
        h = f'Class {self.class_name} - {self.medium} Medium'
        if self._subj:
            h += f' | {self._s(self._subj)}'
        self.cell(0, 6, h, 0, 1, 'C')
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(2)

    def footer(self):
        self.set_y(-15)
        self.set_font('Body', 'I', 8)
        self.set_text_color(160, 160, 160)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', 0, 0, 'C')

    def _s(self, t):
        """Return text as-is, preserving Hindi/Devanagari characters."""
        if not t: return ""
        return str(t)

    def safe_multi_cell(self, w, h, txt):
        # If we are already too close to the right margin, force a new line first
        if self.get_x() > self.w - self.r_margin - 10:
            self.ln()
            
        start_x = self.get_x()
        start_y = self.get_y()
        
        try:
            self.multi_cell(w, h, txt)
        except Exception:
            # A failed multi_cell leaves the cursor at the invalid right margin. Reset it.
            self.set_xy(start_x, start_y)
            
            # Force split on excessively long contiguous characters
            words = txt.split(' ')
            safe_words = []
            for word in words:
                if len(word) > 60:
                    safe_words.append(' '.join(word[i:i+60] for i in range(0, len(word), 60)))
                else:
                    safe_words.append(word)
            try:
                self.multi_cell(w, h, ' '.join(safe_words))
            except Exception:
                # If it still fails, reset again and print fallback text
                self.set_xy(self.l_margin, self.get_y())
                self.multi_cell(w, h, "[Unrenderable content omitted]")

    # ── Cover ──
    def cover(self):
        self.add_page()
        self.ln(50)
        self.set_draw_color(40, 70, 130)
        self.set_line_width(1)
        self.line(50, self.get_y(), 160, self.get_y())
        self.ln(10)
        self.set_font('Body', 'B', 32)
        self.set_text_color(30, 55, 115)
        self.cell(0, 16, f'CLASS {self.class_name}', 0, 1, 'C')
        self.set_font('Body', 'B', 24)
        self.set_text_color(55, 85, 145)
        self.cell(0, 13, f'{self.medium} Medium', 0, 1, 'C')
        self.ln(8)
        self.set_font('Body', '', 15)
        self.set_text_color(80, 80, 80)
        self.cell(0, 9, 'Comprehensive Exam Study Package', 0, 1, 'C')
        self.ln(4)
        self.set_draw_color(40, 70, 130)
        self.set_line_width(0.5)
        self.line(70, self.get_y(), 140, self.get_y())
        self.ln(10)
        self.set_font('Body', '', 11)
        self.set_text_color(100, 100, 100)
        for f in ['Key Concepts & Theorems', 'Important Formulas',
                   'Problem-Solving Methods', 'Exam Focus & Common Pitfalls',
                   'Chapter-wise Structured Notes']:
            self.cell(0, 7, f, 0, 1, 'C')
        self.ln(20)
        self.set_font('Body', 'I', 9)
        self.set_text_color(160, 160, 160)
        self.cell(0, 6, 'NCERT / CBSE Pattern | Exam Oriented', 0, 1, 'C')

    # ── TOC ──
    def toc(self, subjects_data):
        self.add_page()
        self.set_font('Body', 'B', 20)
        self.set_text_color(30, 55, 115)
        self.cell(0, 12, 'TABLE OF CONTENTS', 0, 1, 'C')
        self.ln(6)
        self.set_draw_color(40, 70, 130)
        self.line(60, self.get_y(), 150, self.get_y())
        self.ln(6)
        for subj, subcats in subjects_data.items():
            self.set_font('Body', 'B', 12)
            self.set_text_color(40, 70, 130)
            self.cell(0, 7, self._s(subj), 0, 1)
            for sc_name, chs in subcats.items():
                if sc_name != "_main":
                    self.set_font('Body', 'BI', 10)
                    self.set_text_color(80, 110, 160)
                    self.cell(0, 6, f'  {self._s(sc_name)}', 0, 1)
                self.set_font('Body', '', 9)
                self.set_text_color(80, 80, 80)
                for ch in chs:
                    self.cell(0, 5, f'      {self._s(ch["name"])}', 0, 1)
            self.ln(2)
            if self.get_y() > 260:
                self.add_page()

    # ── Subject divider ──
    def subject_page(self, name):
        self.add_page()
        self._subj = name
        self.ln(70)
        self.set_draw_color(40, 70, 130)
        self.set_line_width(0.8)
        self.line(50, self.get_y(), 160, self.get_y())
        self.ln(8)
        self.set_font('Body', 'B', 28)
        self.set_text_color(30, 55, 115)
        self.cell(0, 14, self._s(name), 0, 1, 'C')
        self.ln(4)
        self.line(60, self.get_y(), 150, self.get_y())

    def subcat_header(self, name):
        self.add_page()
        self.ln(3)
        self.set_font('Body', 'B', 16)
        self.set_text_color(50, 85, 145)
        self.cell(0, 10, self._s(name), 0, 1)
        self.set_draw_color(50, 85, 145)
        self.set_line_width(0.3)
        self.line(10, self.get_y(), 100, self.get_y())
        self.ln(4)

    # ── Render markdown-like GPT output ──
    def render_notes(self, md_text):
        if not md_text:
            return
        for line in md_text.split('\n'):
            s = line.strip()
            if not s:
                self.ln(2)
                continue
            safe = self._s(s)

            # ## Chapter Title
            if s.startswith('## '):
                self.ln(4)
                self.set_font('Body', 'B', 14)
                self.set_text_color(30, 55, 115)
                self.multi_cell(0, 7, safe[3:])
                self.set_draw_color(40, 70, 130)
                self.set_line_width(0.3)
                self.line(10, self.get_y(), 200, self.get_y())
                self.ln(3)
            # ### Section
            elif s.startswith('### '):
                self.ln(3)
                self.set_font('Body', 'B', 12)
                self.set_text_color(50, 80, 140)
                self.safe_multi_cell(0, 6.5, safe[4:])
                self.ln(1.5)
            # #### Sub-section
            elif s.startswith('#### '):
                self.ln(2)
                self.set_font('Body', 'B', 10)
                self.set_text_color(70, 100, 160)
                self.safe_multi_cell(0, 6, safe[5:])
                self.ln(1)
            # Bold line **...**
            elif s.startswith('**') and s.endswith('**'):
                self.set_font('Body', 'B', 10)
                self.set_text_color(40, 40, 40)
                self.safe_multi_cell(0, 5.5, safe.strip('* '))
            # Bullet with bold term  - **Term**: desc
            elif s.startswith('- **') or s.startswith('* **'):
                # Render as bold-prefix bullet
                self.set_font('Body', '', 9.5)
                self.set_text_color(40, 40, 40)
                cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', safe)
                self.safe_multi_cell(0, 5, '  - ' + cleaned[2:].strip(' -*'))
            # Regular bullet
            elif s.startswith('- ') or s.startswith('* '):
                self.set_font('Body', '', 9.5)
                self.set_text_color(40, 40, 40)
                cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', safe)
                self.safe_multi_cell(0, 5, '  - ' + cleaned[2:])
            # Indented bullet
            elif s.startswith('  - ') or s.startswith('  * '):
                self.set_font('Body', '', 9)
                self.set_text_color(60, 60, 60)
                cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', safe)
                self.safe_multi_cell(0, 5, '      ' + cleaned.strip(' -*'))
            # Numbered
            elif re.match(r'^\d+[\.\)]\s', s):
                self.set_font('Body', '', 9.5)
                self.set_text_color(40, 40, 40)
                cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', safe)
                self.safe_multi_cell(0, 5, '  ' + cleaned)
            # LaTeX / formula blocks (render as-is, simplified)
            elif s.startswith('\\[') or s.startswith('$$'):
                self.set_font('Courier', '', 9)
                self.set_text_color(30, 80, 30)
                cleaned = safe.replace('\\[', '').replace('\\]', '').replace('$$', '').strip()
                if cleaned:
                    self.safe_multi_cell(0, 5, '    ' + cleaned)
            # Italic line *...*
            elif s.startswith('*') and s.endswith('*') and not s.startswith('**'):
                self.set_font('Body', 'I', 9)
                self.set_text_color(80, 80, 80)
                self.safe_multi_cell(0, 5, safe.strip('* '))
            # Table rows
            elif s.startswith('|') and s.endswith('|'):
                if '---' in s:
                    continue
                self.set_font('Body', '', 8.5)
                self.set_text_color(40, 40, 40)
                cells = [c.strip() for c in s.split('|') if c.strip()]
                if cells:
                    w = min(180 / max(len(cells), 1), 90)
                    for c in cells:
                        self.cell(w, 5.5, self._s(c)[:55], 1, 0, 'L')
                    self.ln()
            # Regular text
            else:
                self.set_font('Body', '', 9.5)
                self.set_text_color(50, 50, 50)
                cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', safe)
                cleaned = re.sub(r'\*([^*]+)\*', r'\1', cleaned)
                self.safe_multi_cell(0, 5, cleaned)

            if self.get_y() > 275:
                self.add_page()


# ─── Build PDF ─────────────────────────────────────────────────────────────────

def build_pdf(class_name, medium, subjects_data, processed_notes):
    """Assemble final PDF from processed notes."""
    out_file = OUTPUT_DIR / f"Class_{class_name}_{medium}_Medium_Exam_Notes.pdf"

    pdf = ExamPDF(class_name, medium)
    pdf.alias_nb_pages()
    pdf.cover()
    pdf.toc(subjects_data)

    for subj_name, subcats in subjects_data.items():
        pdf.subject_page(subj_name)
        for sc_name, chapters in subcats.items():
            if sc_name != "_main":
                pdf.subcat_header(sc_name)
            for ch in chapters:
                key = f"{subj_name}||{sc_name}||{ch['name']}"
                notes = processed_notes.get(key)
                if notes:
                    pdf.add_page()
                    pdf.render_notes(notes)
                else:
                    # Fallback: render raw text if GPT failed
                    pdf.add_page()
                    pdf.set_font('Body', 'B', 13)
                    pdf.set_text_color(30, 55, 115)
                    pdf.multi_cell(0, 7, pdf._s(ch['name']))
                    pdf.ln(3)
                    pdf.set_font('Body', '', 9)
                    pdf.set_text_color(50, 50, 50)
                    # Render first 3000 chars of raw text as fallback
                    raw = ch['text'][:3000]
                    for ln in raw.split('\n'):
                        sl = ln.strip()
                        if sl:
                            try:
                                pdf.multi_cell(0, 5, pdf._s(sl))
                            except Exception as e:
                                print(f"Skipping unrenderable line in PDF: {e}")
                        else:
                            pdf.ln(2)
                        if pdf.get_y() > 275:
                            pdf.add_page()

    pdf.output(str(out_file))
    sz = out_file.stat().st_size / (1024 * 1024)
    print(f"  >> {out_file.name}  ({sz:.1f} MB, {pdf.page_no()} pages)")
    return out_file


# ─── Main Pipeline ────────────────────────────────────────────────────────────

def process_class(class_name, class_dir, medium):
    label = f"Class {class_name} - {medium} Medium"
    print(f"\n{'='*60}")
    print(f"  {label}")
    print(f"  Source: {class_dir}")
    print(f"{'='*60}")

    if not class_dir.exists():
        print(f"  [SKIP] Not found")
        return None

    # 1. Extract
    print(f"\n  [1/3] Extracting content...")
    subjects = collect_subject_data(class_dir)
    if not subjects:
        print(f"  [SKIP] No content")
        return None

    total = sum(len(c) for sc in subjects.values() for c in sc.values())
    print(f"  Found {len(subjects)} subjects, {total} chapters\n")

    # 2. Process through GPT
    print(f"  [2/3] Generating structured notes via GPT...")
    
    cache_file = Path(f"rag_cache_Class_{class_name}_{medium}.json")
    if cache_file.exists():
        with open(cache_file, "r", encoding="utf-8") as f:
            processed = json.load(f)
    else:
        processed = {}
        
    done = 0
    for subj_name, subcats in subjects.items():
        for sc_name, chapters in subcats.items():
            sc_display = f"{subj_name}" + (f"/{sc_name}" if sc_name != "_main" else "")
            for ch in chapters:
                done += 1
                key = f"{subj_name}||{sc_name}||{ch['name']}"
                print(f"    [{done}/{total}] {sc_display} > {ch['name'][:50]}...", end=" ", flush=True)

                if key in processed:
                    print("SKIPPED (Cached)")
                    continue

                notes = process_chapter_gpt(ch['name'], ch['text'], subj_name, f"Class {class_name}", medium)
                if notes:
                    processed[key] = notes
                    # Save to cache immediately
                    with open(cache_file, "w", encoding="utf-8") as f:
                        json.dump(processed, f, ensure_ascii=False, indent=2)
                    print("OK")
                else:
                    print("FAILED (raw fallback)")

                # Small delay to avoid rate limits
                time.sleep(0.5)

    print(f"\n  Processed {len(processed)}/{total} chapters with AI")

    # 3. Build PDF
    print(f"\n  [3/3] Building PDF...")
    return build_pdf(class_name, medium, subjects, processed)


def main():
    global client, OPENAI_API_KEY

    print("=" * 60)
    print("  EXAM MATERIAL PDF GENERATOR")
    print("  Classes 6-8 | English & Hindi Medium")
    print("=" * 60)

    OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
    if not OPENAI_API_KEY:
        print("\n[ERROR] Set your API key first:")
        print('  $env:OPENAI_API_KEY = "sk-..."')
        sys.exit(1)

    # Detect API keys and set base URL accordingly
    if OPENAI_API_KEY.startswith("sk-or-"):
        client = openai.OpenAI(
            api_key=OPENAI_API_KEY,
            base_url="https://openrouter.ai/api/v1"
        )
        print("  Using: OpenRouter API")
    elif OPENAI_API_KEY.startswith("AIza"):
        client = openai.OpenAI(
            api_key=OPENAI_API_KEY,
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
        )
        print("  Using: Gemini API")
    elif OPENAI_API_KEY.lower() == "ollama":
        client = openai.OpenAI(
            api_key="ollama",
            base_url="http://localhost:11434/v1"
        )
        print("  Using: Local Ollama API")
    else:
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        print("  Using: OpenAI API")
    print(f"\nOutput: {OUTPUT_DIR}\n")

    results = []

    # English Medium (temporarily disabled per request)
    # for cls in ["6th", "7th", "8th"]:
    #     r = process_class(cls, ENGLISH_MEDIUM[cls], "English")
    #     if r: results.append(r)

    # Hindi Medium
    for cls in HINDI_MEDIUM.keys():
        r = process_class(cls, HINDI_MEDIUM[cls], "Hindi")
        if r: results.append(r)

    print(f"\n{'='*60}")
    print(f"  DONE! {len(results)} PDFs generated:")
    print(f"{'='*60}")
    for r in results:
        print(f"  {r.name}  ({r.stat().st_size/(1024*1024):.1f} MB)")
    print(f"\n  Location: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()

